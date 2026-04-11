# app/frames/transcriber.py
import os
import sys
import queue
import threading
import time
import tkinter as tk
from tkinter import filedialog, messagebox

import ttkbootstrap as ttk
from ttkbootstrap.constants import *

import time
import subprocess

from app.utils import HAS_DND, HAS_FW, HAS_DOCX, DND_FILES, _ext, seconds_to_hms, create_no_window_flags, ffprobe_cmd

# ---------- Transcrição ----------
AUDIO_VIDEO_EXTS = {"mp3", "wav", "m4a", "mp4", "mkv", "mov", "webm"}
SUPPORTED_EXTS = AUDIO_VIDEO_EXTS

# Modelo atual (você pode trocar para "medium" se quiser mais qualidade)
WHISPER_MODEL = "large-v2"

# Ajustes para transcrição mais completa
PRIMARY_BEAM_SIZE = 10          # mais completo (mais lento)
FALLBACK_BEAM_SIZE = 16         # fallback ainda mais completo (mais lento)
MIN_TEXT_CHARS_FOR_OK = 120     # se sair menos que isso, considera fraco e tenta fallback
HEARTBEAT_IDLE_SECONDS = 3.0


class TranscriberFrame(ttk.Frame):
    def __init__(self, master, on_status):
        super().__init__(master)
        self.on_status = on_status

        self.input_files = []
        self.progress_var = tk.DoubleVar(value=0)
        self.status_var = tk.StringVar(value="")
        self.output_name_var = tk.StringVar()

        self.is_running = False
        self.cancel_requested = False
        self.ui_queue = queue.Queue()

        self.last_output = None
        self.model = None
        self._completed_file_times = []
        self._last_progress_ts = 0.0
        self._heartbeat_running = False
        self._last_action_key_ts = 0.0

        self._build_ui()
        self.bind_all("<Return>", self._handle_return_key, add="+")
        self.bind_all("<Escape>", self._handle_escape_key, add="+")
        self.after(100, self._drain_ui_queue)

        if HAS_DND:
            try:
                self.drop_target_register(DND_FILES)
                self.dnd_bind("<<Drop>>", self._on_drop_files)
            except Exception:
                pass

    def _build_ui(self):
        card = ttk.Frame(self, padding=18)
        card.pack(fill="both", expand=True)

        header = ttk.Frame(card)
        header.pack(fill="x")
        ttk.Label(header, text="Transcritor de Audio (Word)", style="SectionTitle.TLabel").pack(side="left")
        ttk.Separator(card).pack(fill="x", pady=12)

        files_frame = ttk.LabelFrame(card, text="Arquivos")
        files_frame.pack(fill="x")
        files_inner = ttk.Frame(files_frame, padding=12)
        files_inner.pack(fill="x")
        files_inner.columnconfigure(1, weight=1)

        ttk.Button(files_inner, text="Selecionar Arquivo(s)", command=self.selecionar_arquivos, bootstyle=WARNING).grid(
            row=0, column=0, sticky="w"
        )
        ttk.Button(files_inner, text="Remover", command=self.remover_arquivos, bootstyle=DANGER).grid(
            row=0, column=1, sticky="w", padx=(10, 0)
        )

        self.label_sel = ttk.Label(files_inner, text="Nenhum arquivo selecionado", font=("Helvetica", 12))
        self.label_sel.grid(row=1, column=0, columnspan=2, sticky="w", pady=(10, 0))

        if HAS_DND:
            ttk.Label(
                files_inner,
                text="Arraste e solte audio/video aqui para selecionar.",
                style="Muted.TLabel",
            ).grid(row=2, column=0, columnspan=2, sticky="w", pady=(6, 4))

        self.opts_frame = ttk.LabelFrame(card, text="Saída")
        self.opts_frame.pack_forget()
        output_inner = ttk.Frame(self.opts_frame, padding=12)
        output_inner.pack(fill="x")
        output_inner.columnconfigure(1, weight=1)

        ttk.Label(output_inner, text="Nome do arquivo:", font=("Helvetica", 13, "bold")).grid(row=0, column=0, sticky="w")
        self.output_name_entry = ttk.Entry(output_inner, textvariable=self.output_name_var, width=32)
        self.output_name_entry.grid(row=0, column=1, sticky="ew", padx=(10, 0))
        ttk.Label(output_inner, text="Editavel quando houver 1 arquivo selecionado.", style="Muted.TLabel").grid(
            row=0, column=2, sticky="w", padx=(10, 0)
        )
        self.output_name_entry.configure(state=DISABLED)

        self.controls_frame = ttk.Frame(card)
        self.controls_frame.pack_forget()
        self.btn_run = ttk.Button(self.controls_frame, text="Transcrever (mais completo)", command=self.start_transcription, bootstyle=SUCCESS, state=DISABLED)
        self.btn_run.pack(side="left")
        self.btn_cancel = ttk.Button(self.controls_frame, text="Cancelar", command=self.cancel_transcription, bootstyle=SECONDARY, state=DISABLED)
        self.btn_cancel.pack(side="left", padx=(10, 0))

        self.progress_frame = ttk.Frame(card, padding=10)
        self.progress = ttk.Progressbar(self.progress_frame, orient=tk.HORIZONTAL, mode="determinate", variable=self.progress_var, maximum=100)
        self.progress.pack(fill="x")
        self.status_lbl = ttk.Label(self.progress_frame, textvariable=self.status_var, font=("Helvetica", 11))
        self.status_lbl.pack(anchor="w", pady=(6, 0))

        # Hide progress UI until transcription starts
        self._hide_progress()

        self.btn_open = ttk.Button(card, text="Abrir pasta do ultimo .docx", command=self.abrir_pasta, bootstyle=INFO, state=DISABLED)
        self.btn_open.pack_forget()
        self._update_action_state()
        self._update_visibility()

    def selecionar_arquivos(self):
        tipos = [
            ("Audio/Video", "*.mp3 *.wav *.m4a *.mp4 *.mkv *.mov *.webm"),
            ("Audio", "*.mp3 *.wav *.m4a"),
            ("Video", "*.mp4 *.mkv *.mov *.webm"),
            ("Todos", "*.*"),
        ]
        paths = filedialog.askopenfilenames(title="Selecione arquivo(s)", filetypes=tipos)
        if paths:
            self._set_files(list(paths))

    def _on_drop_files(self, event):
        items = self.tk.splitlist(event.data)
        paths = []

        for p in items:
            if os.path.isdir(p):
                try:
                    for nm in os.listdir(p):
                        f = os.path.join(p, nm)
                        if os.path.isfile(f) and _ext(f) in SUPPORTED_EXTS:
                            paths.append(os.path.abspath(f))
                except Exception:
                    pass
            elif os.path.isfile(p) and _ext(p) in SUPPORTED_EXTS:
                paths.append(os.path.abspath(p))

        seen = set()
        uniq = []
        for p in paths:
            low = p.lower()
            if low not in seen:
                seen.add(low)
                uniq.append(p)

        if uniq:
            self._set_files(uniq)

    def _set_files(self, paths):
        paths = [p for p in paths if os.path.isfile(p) and _ext(p) in SUPPORTED_EXTS]
        self.input_files = paths

        if not self.input_files:
            self.label_sel.config(text="Nenhum arquivo selecionado")
            self._refresh_output_name()
            self._update_action_state()
            return

        self.label_sel.config(
            text=f"Arquivo: {os.path.basename(paths[0])}"
            if len(paths) == 1 else f"{len(paths)} arquivos (ex.: {os.path.basename(paths[0])})"
        )
        self._refresh_output_name()
        self._update_action_state()

    def _refresh_output_name(self):
        if not self.input_files:
            self.output_name_var.set("")
            self.output_name_entry.configure(state=DISABLED)
            return

        if len(self.input_files) > 1:
            self.output_name_var.set("")
            self.output_name_entry.configure(state=DISABLED)
            return

        filename = os.path.splitext(os.path.basename(self.input_files[0]))[0]
        self.output_name_var.set(f"{filename}.docx")
        self.output_name_entry.configure(state=NORMAL)

    def _build_output_path(self, in_path):
        if len(self.input_files) == 1:
            custom_name = (self.output_name_var.get() or "").strip()
            if custom_name:
                if "." not in os.path.basename(custom_name):
                    custom_name += ".docx"
                return os.path.join(os.path.dirname(in_path), custom_name)
        return os.path.join(
            os.path.dirname(in_path),
            os.path.splitext(os.path.basename(in_path))[0] + ".docx"
        )

    def _is_active_screen(self):
        top = self.winfo_toplevel()
        return getattr(top, "current_screen", None) == getattr(self, "screen_key", None)

    def _handle_return_key(self, event=None):
        if not self._is_active_screen() or self.is_running or str(self.btn_run["state"]) != str(NORMAL):
            return
        if isinstance(event.widget, tk.Text):
            return
        now = time.monotonic()
        if now - self._last_action_key_ts < 0.35:
            return "break"
        self._last_action_key_ts = now
        self.start_transcription()
        return "break"

    def _handle_escape_key(self, _event=None):
        if not self._is_active_screen() or not self.is_running:
            return
        self.cancel_transcription()
        return "break"

    def _update_action_state(self):
        if self.is_running:
            self.btn_run.config(state=DISABLED)
            self.btn_cancel.config(state=NORMAL)
            self._update_visibility()
            return

        self.btn_run.config(state=NORMAL if self.input_files else DISABLED)
        self.btn_cancel.config(state=DISABLED)
        self._update_visibility()

    def _update_visibility(self):
        """Show/hide output and control panels when there are selected files."""
        if self.input_files:
            if not self.opts_frame.winfo_ismapped():
                self.opts_frame.pack(fill="x", pady=(10, 0))
            if not self.controls_frame.winfo_ismapped():
                self.controls_frame.pack(fill="x", pady=(10, 6))
            if not self.btn_open.winfo_ismapped():
                self.btn_open.pack(pady=8)
        else:
            self.opts_frame.pack_forget()
            self.controls_frame.pack_forget()
            self.btn_open.pack_forget()

        # Only show progress when transcription is running.
        if self.is_running:
            self._show_progress()
        else:
            self._hide_progress()

    def _show_progress(self):
        if getattr(self, "progress_frame", None) and not self.progress_frame.winfo_ismapped():
            self.progress_frame.pack(fill="x", pady=(8, 4))

    def _hide_progress(self):
        if getattr(self, "progress_frame", None) and self.progress_frame.winfo_ismapped():
            self.progress_frame.pack_forget()

    def remover_arquivos(self):
        self.input_files = []
        self.label_sel.config(text="Nenhum arquivo selecionado")
        self.progress_var.set(0)
        self.status_var.set("")
        self._hide_progress()
        self.btn_open.config(state=DISABLED)
        self.output_name_var.set("")
        self.output_name_entry.configure(state=DISABLED)
        self._update_action_state()

    def start_transcription(self):
        if self.is_running:
            return

        if not HAS_FW or not HAS_DOCX:
            messagebox.showerror("Dependências", "Para transcrever instale: pip install faster-whisper python-docx")
            return

        if not self.input_files:
            messagebox.showerror("Erro", "Selecione pelo menos um arquivo.")
            return

        if len(self.input_files) == 1:
            custom_name = (self.output_name_var.get() or "").strip()
            if not custom_name:
                messagebox.showerror("Erro", "Informe um nome para o arquivo de saida.")
                return
            proposed_output = self._build_output_path(self.input_files[0])
            if os.path.abspath(proposed_output).lower() == os.path.abspath(self.input_files[0]).lower():
                messagebox.showerror("Erro", "Escolha um nome diferente do arquivo original.")
                return

        self.is_running = True
        self.cancel_requested = False

        self.btn_run.config(state=DISABLED)
        self.btn_cancel.config(state=NORMAL)
        self.btn_open.config(state=DISABLED)
        self._show_progress()
        self._update_action_state()

        self.progress_var.set(0)
        self.status_var.set("Preparando...")
        self.on_status("Transcrição iniciada…")
        self._last_progress_ts = time.time()
        self._set_progress_mode("determinate")

        threading.Thread(target=self._batch_transcribe_worker, args=(self.input_files[:],), daemon=True).start()

    def cancel_transcription(self):
        if self.is_running:
            self.cancel_requested = True

    def _set_progress_mode(self, mode):
        if mode == "indeterminate":
            if not self._heartbeat_running:
                self.progress.configure(mode="indeterminate")
                self.progress.start(10)
                self._heartbeat_running = True
            return

        if self._heartbeat_running:
            self.progress.stop()
            self._heartbeat_running = False
        self.progress.configure(mode="determinate")

    def _update_heartbeat_state(self):
        if not self.is_running:
            self._set_progress_mode("determinate")
            return
        idle = time.time() - self._last_progress_ts
        if idle >= HEARTBEAT_IDLE_SECONDS:
            self._set_progress_mode("indeterminate")
        else:
            self._set_progress_mode("determinate")

    def _ensure_model(self):
        if self.model is not None:
            return
        self.ui_queue.put(("status", f"Carregando modelo '{WHISPER_MODEL}'..."))
        from faster_whisper import WhisperModel # type: ignore
        # CPU int8: bom pra uso geral em PC comum
        self.model = WhisperModel(WHISPER_MODEL, device="cpu", compute_type="int8")
        self.ui_queue.put(("status", "Modelo carregado."))

    def _batch_transcribe_worker(self, files):
        try:
            self._ensure_model()
            total = len(files)
            last_out = None
            successes = 0
            failures = 0

            for idx, path in enumerate(files, start=1):
                if self.cancel_requested:
                    break

                out_path = self._build_output_path(path)

                self.ui_queue.put(("status", f"[{idx}/{total}] Transcrevendo: {os.path.basename(path)}"))
                ok = self._transcribe_one(path, out_path, idx, total)

                if ok:
                    last_out = out_path
                    successes += 1
                else:
                    failures += 1

            if self.cancel_requested:
                self.ui_queue.put(("canceled", "Transcrição cancelada."))
            else:
                if failures:
                    message = f"Transcrição concluída: {successes} de {total} arquivo(s). {failures} falharam."
                else:
                    message = f"Transcrição concluída de {successes} arquivo(s)."
                self.ui_queue.put((
                    "done",
                    {
                        "message": message,
                        "successes": successes,
                        "failures": failures,
                        "total": total,
                        "last_output": last_out,
                    },
                ))

        except Exception as e:
            if self.cancel_requested:
                self.ui_queue.put(("canceled", "Transcrição cancelada."))
            else:
                self.ui_queue.put(("error", f"Erro no processamento: {e}"))
        finally:
            self.is_running = False

    def _save_docx(self, text, out_docx, info):
        try:
            from docx import Document # type: ignore
        except Exception as e:
            raise RuntimeError("Dependências ausentes: instale python-docx (pip install python-docx).") from e

        doc = Document()
        doc.add_heading("Transcrição", level=1)

        # Idioma detectado
        lang = getattr(info, "language", None)
        if lang:
            doc.add_paragraph(f"Idioma detectado: {lang}")
            doc.add_paragraph("")

        doc.add_paragraph(text)
        doc.save(out_docx)
        return True

    def _run_transcribe(self, in_path: str, beam_size: int):
        """
        Executa transcribe com foco em completude:
        - language=None (multi-idioma)
        - vad_filter=False (não cortar fala)
        - beam_size configurável (mais completo)
        """
        return self.model.transcribe(
            in_path,
            language=None,
            vad_filter=False,
            beam_size=beam_size,
            condition_on_previous_text=True
        )

    def _probe_duration(self, path):
        try:
            out = subprocess.check_output(
                ffprobe_cmd("-v", "error", "-show_entries", "format=duration", "-of", "default=nk=1:nw=1", path),
                text=True,
                creationflags=create_no_window_flags(),
                stderr=subprocess.DEVNULL,
            )
            try:
                return float(out.strip())
            except Exception:
                return None
        except Exception:
            return None

    def _transcribe_one(self, in_path, out_docx, idx, total):
        try:
            file_start = time.time()

            # 1) Primeira tentativa (mais completa que o padrão)
            # Probe duration to help ETA estimation
            duration = self._probe_duration(in_path)

            # Run transcribe in a background thread and poll so we can provide ETA/progress
            model_result = {}
            def _run_primary():
                try:
                    model_result['value'] = self._run_transcribe(in_path, beam_size=PRIMARY_BEAM_SIZE)
                except Exception as e:
                    model_result['error'] = e

            t = threading.Thread(target=_run_primary, daemon=True)
            t.start()

            pieces = []
            # Poll while the model is running to show ETA even if segments are not yielded incrementally
            poll_start = time.time()
            while t.is_alive():
                if self.cancel_requested:
                    break
                elapsed = time.time() - poll_start

                # estimate expected time for this file
                if self._completed_file_times:
                    avg_file_time = sum(self._completed_file_times) / len(self._completed_file_times)
                elif duration:
                    avg_file_time = max(1.0, duration * 0.6)
                else:
                    avg_file_time = max(5.0, elapsed * 2)

                per_file_frac = elapsed / (elapsed + avg_file_time)
                overall_frac = ((idx - 1) + per_file_frac) / total
                pct = max(0.0, min(100.0, overall_frac * 100.0))

                overall_remaining = max(0.0, avg_file_time - elapsed) + (total - idx) * avg_file_time
                eta = seconds_to_hms(overall_remaining)

                try:
                    self.ui_queue.put(("progress", pct))
                    self.ui_queue.put(("status", f"[{idx}/{total}] Transcrevendo... {pct:.1f}% — {eta} restante"))
                except Exception:
                    pass

                time.sleep(0.6)

            if self.cancel_requested:
                return False

            # thread finished
            if 'error' in model_result:
                raise model_result['error']

            segments, info = model_result.get('value', ([], None))
            for i, seg in enumerate(segments, start=1):
                if self.cancel_requested:
                    break
                t = (seg.text or "").strip()
                if t:
                    pieces.append(t)

                try:
                    if duration:
                        seg_end = getattr(seg, "end", None)
                        if seg_end is not None:
                            per_file_frac = max(0.0, min(1.0, float(seg_end) / max(duration, 1e-9)))
                        else:
                            elapsed_full = time.time() - poll_start
                            avg_file_time = max(1.0, duration * 0.6)
                            per_file_frac = elapsed_full / (elapsed_full + avg_file_time)
                    else:
                        elapsed_full = time.time() - poll_start
                        if self._completed_file_times:
                            avg_file_time = sum(self._completed_file_times) / len(self._completed_file_times)
                        else:
                            avg_file_time = max(5.0, elapsed_full * 2)
                        per_file_frac = elapsed_full / (elapsed_full + avg_file_time)

                    overall_frac = ((idx - 1) + per_file_frac) / total
                    pct = max(0.0, min(100.0, overall_frac * 100.0))

                    elapsed_full = time.time() - poll_start
                    if per_file_frac > 0:
                        expected_total = elapsed_full / per_file_frac
                    else:
                        expected_total = elapsed_full
                    remaining_curr = max(0.0, expected_total - elapsed_full)

                    if self._completed_file_times:
                        avg_file_time = sum(self._completed_file_times) / len(self._completed_file_times)
                    else:
                        avg_file_time = elapsed_full + remaining_curr

                    overall_remaining = remaining_curr + (total - idx) * avg_file_time
                    eta = seconds_to_hms(overall_remaining)

                    self.ui_queue.put(("progress", pct))
                    self.ui_queue.put(("status", f"[{idx}/{total}] Transcrevendo... {pct:.1f}% — {eta} restante"))
                except Exception:
                    pass

            if self.cancel_requested:
                return False

            full = " ".join(pieces).strip()

            # 2) Fallback automático: se veio pouco texto, tentar ainda mais completo
            if len(full) < MIN_TEXT_CHARS_FOR_OK:
                self.ui_queue.put(("status", f"[{idx}/{total}] Resultado curto — tentando modo mais completo..."))

                # Run fallback transcribe in background and poll
                fb_result = {}
                def _run_fallback():
                    try:
                        fb_result['value'] = self._run_transcribe(in_path, beam_size=FALLBACK_BEAM_SIZE)
                    except Exception as e:
                        fb_result['error'] = e

                tfb = threading.Thread(target=_run_fallback, daemon=True)
                tfb.start()

                fb_start = time.time()
                while tfb.is_alive():
                    if self.cancel_requested:
                        break
                    elapsed = time.time() - fb_start

                    if self._completed_file_times:
                        avg_file_time = sum(self._completed_file_times) / len(self._completed_file_times)
                    elif duration:
                        avg_file_time = max(1.0, duration * 0.6)
                    else:
                        avg_file_time = max(5.0, elapsed * 2)

                    per_file_frac = elapsed / (elapsed + avg_file_time)
                    overall_frac = ((idx - 1) + per_file_frac) / total
                    pct = max(0.0, min(100.0, overall_frac * 100.0))

                    overall_remaining = max(0.0, avg_file_time - elapsed) + (total - idx) * avg_file_time
                    eta = seconds_to_hms(overall_remaining)

                    try:
                        self.ui_queue.put(("progress", pct))
                        self.ui_queue.put(("status", f"[{idx}/{total}] Tentando modo mais completo... {pct:.1f}% — {eta} restante"))
                    except Exception:
                        pass

                    time.sleep(0.6)

                if self.cancel_requested:
                    return False

                if 'error' in fb_result:
                    raise fb_result['error']

                segments2, info2 = fb_result.get('value', ([], None))
                pieces2 = []
                for j, seg in enumerate(segments2, start=1):
                    if self.cancel_requested:
                        break
                    t = (seg.text or "").strip()
                    if t:
                        pieces2.append(t)

                    # update progress during fallback pass + ETA
                    try:
                        elapsed_full = time.time() - fb_start
                        if duration:
                            seg_end = getattr(seg, "end", None)
                            if seg_end is not None:
                                per_file_frac = max(0.0, min(1.0, float(seg_end) / max(duration, 1e-9)))
                            else:
                                avg_file_time = max(1.0, duration * 0.6)
                                per_file_frac = elapsed_full / (elapsed_full + avg_file_time)
                        else:
                            if self._completed_file_times:
                                avg_file_time = sum(self._completed_file_times) / len(self._completed_file_times)
                            else:
                                avg_file_time = max(5.0, elapsed_full * 2)
                            per_file_frac = elapsed_full / (elapsed_full + avg_file_time)

                        overall_frac = ((idx - 1) + per_file_frac) / total
                        pct = max(0.0, min(100.0, overall_frac * 100.0))

                        if per_file_frac > 0:
                            expected_total = elapsed_full / per_file_frac
                        else:
                            expected_total = elapsed_full
                        remaining_curr = max(0.0, expected_total - elapsed_full)

                        if self._completed_file_times:
                            avg_file_time = sum(self._completed_file_times) / len(self._completed_file_times)
                        else:
                            avg_file_time = elapsed_full + remaining_curr

                        overall_remaining = remaining_curr + (total - idx) * avg_file_time
                        eta = seconds_to_hms(overall_remaining)

                        self.ui_queue.put(("progress", pct))
                        self.ui_queue.put(("status", f"[{idx}/{total}] Tentando modo mais completo... {pct:.1f}% — {eta} restante"))
                    except Exception:
                        pass

                if self.cancel_requested:
                    return False

                full2 = " ".join(pieces2).strip()

                # usa o melhor (mais longo)
                if len(full2) > len(full):
                    full = full2
                    info = info2

            if not full:
                self.ui_queue.put(("error", f"Nenhum texto reconhecido em: {os.path.basename(in_path)}"))
                return False

            try:
                self._save_docx(full, out_docx, info)
            except RuntimeError as e:
                self.ui_queue.put(("error", str(e)))
                return False

            # record elapsed time for this file
            try:
                file_elapsed = time.time() - file_start
                self._completed_file_times.append(file_elapsed)
            except Exception:
                pass

            # mark this file as completed (progress reaches fraction for this file)
            try:
                pct_done = (idx / total) * 100.0
                self.ui_queue.put(("progress", pct_done))
                self.ui_queue.put(("status", f"[{idx}/{total}] Salvo: {os.path.basename(out_docx)}"))
            except Exception:
                pass

            return True

        except Exception as e:
            self.ui_queue.put(("error", f"Erro ao transcrever {os.path.basename(in_path)}: {e}"))
            return False

    def _drain_ui_queue(self):
        try:
            while True:
                kind, payload = self.ui_queue.get_nowait()

                if kind == "progress":
                    self.progress_var.set(payload)
                    self._last_progress_ts = time.time()
                    self._set_progress_mode("determinate")

                elif kind == "status":
                    self.status_var.set(payload)
                    self.on_status(payload)

                elif kind == "done":
                    info = payload if isinstance(payload, dict) else {"message": str(payload)}
                    message = info.get("message") or "Transcrição concluída."
                    successes = int(info.get("successes", 0) or 0)
                    failures = int(info.get("failures", 0) or 0)
                    last_output = info.get("last_output")

                    if last_output:
                        self.last_output = last_output

                    self._hide_progress()
                    self.progress_var.set(100 if successes > 0 else 0)
                    self.status_var.set(message)
                    self.on_status(message)

                    self._update_action_state()
                    self.btn_open.config(state=NORMAL if self.last_output else DISABLED)
                    self._update_visibility()
                    self._set_progress_mode("determinate")

                    if failures:
                        messagebox.showwarning("Aviso", message)
                    elif successes > 0:
                        messagebox.showinfo("Sucesso", message)
                    else:
                        messagebox.showerror("Erro", message)

                elif kind == "canceled":
                    self._hide_progress()
                    self.status_var.set(payload)
                    self.on_status(payload)
                    self.progress_var.set(0)

                    self._update_action_state()
                    self.btn_open.config(state=DISABLED)
                    self._update_visibility()
                    self._set_progress_mode("determinate")

                elif kind == "error":
                    self._hide_progress()
                    self.on_status("Erro na transcrição")
                    messagebox.showerror("Erro", str(payload))
                    self._set_progress_mode("determinate")

        except queue.Empty:
            pass
        finally:
            self._update_heartbeat_state()
            if not self.is_running and self.btn_cancel["state"] == NORMAL:
                self._update_action_state()
            if self.winfo_exists():
                self.after(100, self._drain_ui_queue)

    def abrir_pasta(self):
        if self.last_output and os.path.exists(self.last_output):
            pasta = os.path.dirname(self.last_output)
            try:
                if sys.platform.startswith("win"):
                    os.startfile(pasta)
                elif sys.platform == "darwin":
                    import subprocess
                    subprocess.Popen(["open", pasta])
                else:
                    import subprocess
                    subprocess.Popen(["xdg-open", pasta])
            except Exception as e:
                messagebox.showerror("Erro", f"Não foi possível abrir a pasta: {e}")
        else:
            messagebox.showerror("Erro", "Nenhum arquivo transcrito encontrado.")
