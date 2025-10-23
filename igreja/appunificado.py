# app_unificado_pretty.py
# Mídia Suite: Conversor • YouTube • Transcrição — tema "darkly"
# Interface modernizada com barra lateral, cabeçalho, cartões e statusbar
# Mantém a mesma lógica funcional dos seus três apps originais

import threading
import tkinter as tk
from tkinter import filedialog, messagebox
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import os
import sys
import subprocess
import queue
import json
import math
from pathlib import Path

# =======================
# Dependências opcionais
# =======================
# Drag & Drop
try:
    from tkinterdnd2 import DND_FILES, TkinterDnD
    HAS_DND = True
except Exception:
    HAS_DND = False

# Imagens (conversor)
try:
    from PIL import Image
    HAS_PIL = True
except Exception:
    HAS_PIL = False

try:
    import rawpy
    HAS_RAWPY = True
except Exception:
    HAS_RAWPY = False

# Transcrição (faster-whisper + docx)
_MISSING_TRANSCRIBE_DEPS = []
try:
    from faster_whisper import WhisperModel
    HAS_FW = True
except Exception:
    HAS_FW = False
    _MISSING_TRANSCRIBE_DEPS.append("faster-whisper")

try:
    from docx import Document
    HAS_DOCX = True
except Exception:
    HAS_DOCX = False
    _MISSING_TRANSCRIBE_DEPS.append("python-docx")

# =======================
# Utils gerais
# =======================
def create_no_window_flags():
    if sys.platform.startswith("win"):
        return subprocess.CREATE_NO_WINDOW
    return 0

def seconds_to_hms(s):
    try:
        s = float(s)
    except Exception:
        return "00:00:00"
    h = int(s // 3600)
    m = int((s % 3600) // 60)
    sec = int(s % 60)
    return f"{h:02d}:{m:02d}:{sec:02d}"

def _ext(path):
    return os.path.splitext(path)[1][1:].lower()

def format_bytes(n):
    try:
        n = float(n)
    except Exception:
        return None
    if n <= 0:
        return "0 B"
    units = ["B", "KB", "MB", "GB", "TB"]
    idx = min(int(math.log(n, 1024)), len(units) - 1)
    val = n / (1024 ** idx)
    return (f"{val:,.0f}" if val >= 100 else f"{val:,.2f}").replace(",", ".") + f" {units[idx]}"

# ======================= Conversor =======================
VIDEO_EXTS = {"mp4", "avi", "mkv", "mov", "webm", "flv", "m4v"}
AUDIO_EXTS = {"wav", "mp3"}
IMAGE_EXTS = {"jpg", "jpeg", "png", "webp", "bmp", "tif", "tiff", "cr2"}
ALL_EXTS = VIDEO_EXTS | AUDIO_EXTS | IMAGE_EXTS

def is_video_file(path):
    e = _ext(path)
    return e in VIDEO_EXTS or e in AUDIO_EXTS

def is_image_file(path):
    return _ext(path) in IMAGE_EXTS

class ConverterFrame(ttk.Frame):
    def __init__(self, master, on_status):
        super().__init__(master)
        self.on_status = on_status  # callback para statusbar global

        self.input_files = []
        self.ultimo_arquivo_convertido = ""
        self.formato_destino = tk.StringVar(value="mp4")
        self.progress_var = tk.DoubleVar(value=0)
        self.status_var = tk.StringVar(value="")
        self.is_converting = False
        self.proc = None
        self.cancel_requested = False
        self._current_output_path = None
        self.ui_queue = queue.Queue()

        self.video_formats = ["mp3", "mp4", "avi", "mkv", "mov"]
        self.image_formats = ["jpg", "png", "webp", "tiff", "bmp"]
        self.current_mode = "video"

        self._build_ui()
        self.after(100, self._drain_ui_queue)

        if HAS_DND:
            try:
                self.drop_target_register(DND_FILES)
                self.dnd_bind("<<Drop>>", self._on_drop_files)
            except Exception:
                pass

    def _build_ui(self):
        # Card wrapper
        card = ttk.Frame(self, padding=18, bootstyle="dark")
        card.pack(fill="both", expand=True)

        header = ttk.Frame(card)
        header.pack(fill="x")
        ttk.Label(header, text="⚙️  Conversor de Vídeo / Imagem", font=("Helvetica", 18, "bold")).pack(side="left")
        ttk.Separator(card).pack(fill="x", pady=12)

        # Linha de ações
        actions = ttk.Frame(card)
        actions.pack(fill="x")
        ttk.Button(actions, text="Selecionar Arquivo(s)", command=self.selecionar_arquivos, bootstyle=WARNING).pack(side="left")
        ttk.Button(actions, text="🗑 Remover", command=self.remover_arquivos, bootstyle=DANGER).pack(side="left", padx=(10, 0))

        # Info seleção
        info = ttk.Frame(card)
        info.pack(fill="x", pady=(10, 0))
        self.label_video = ttk.Label(info, text="Nenhum arquivo selecionado", font=("Helvetica", 12))
        self.label_video.pack(anchor="w")
        self.label_formato = ttk.Label(info, text="", font=("Helvetica", 12))
        self.label_formato.pack(anchor="w", pady=(2,0))

        if HAS_DND:
            ttk.Label(card, text="Dica: arraste e solte vídeos/áudios (mp4, mkv, mp3...) ou imagens (jpg, png, cr2...)", font=("Helvetica", 10, "italic"), foreground="#9aa0a6").pack(anchor="w", pady=(6, 4))

        # Linha formato
        self.fmt_row = ttk.Frame(card)
        self.fmt_row_visible = False
        ttk.Label(self.fmt_row, text="Converter para:", font=("Helvetica", 13, "bold")).pack(side="left")
        self.format_menu = ttk.Combobox(self.fmt_row, textvariable=self.formato_destino, values=[], state="readonly", width=14)
        self.format_menu.pack(side="left", padx=(10, 0))

        # Botões Converter/Cancelar
        ctl = ttk.Frame(card)
        ctl.pack(fill="x", pady=(10, 6))
        self.convert_btn = ttk.Button(ctl, text="▶️ Converter", command=self.start_conversion, bootstyle=SUCCESS)
        self.convert_btn.pack(side="left")
        self.cancel_btn = ttk.Button(ctl, text="Cancelar", command=self.cancel_conversion, bootstyle=SECONDARY, state=DISABLED)
        self.cancel_btn.pack(side="left", padx=(10, 0))

        # Progresso + status
        prog_card = ttk.Frame(card, padding=10, bootstyle="secondary")
        prog_card.pack(fill="x", pady=(8, 4))
        self.progress = ttk.Progressbar(prog_card, orient=tk.HORIZONTAL, mode="determinate", variable=self.progress_var, maximum=100)
        self.progress.pack(fill="x")
        self.status_label = ttk.Label(prog_card, textvariable=self.status_var, font=("Helvetica", 11))
        self.status_label.pack(anchor="w", pady=(6,0))

        # Abrir pasta
        self.open_btn = ttk.Button(card, text="Abrir pasta do arquivo convertido", command=self.abrir_pasta, bootstyle=INFO, state=DISABLED)
        self.open_btn.pack(pady=8)

    # visibilidade combobox
    def _show_format_row(self):
        if not self.fmt_row_visible:
            self.fmt_row.pack(fill="x", pady=(8, 5))
            self.fmt_row_visible = True

    def _hide_format_row(self):
        if self.fmt_row_visible:
            self.fmt_row.pack_forget()
            self.fmt_row_visible = False

    def _update_format_menu(self):
        if not self.input_files:
            self._hide_format_row(); self.format_menu.config(values=[]); return
        all_video = all(is_video_file(p) for p in self.input_files)
        all_image = all(is_image_file(p) for p in self.input_files)
        if not (all_video or all_image):
            messagebox.showerror("Seleção inválida", "Selecione apenas arquivos de VÍDEO/ÁUDIO ou apenas IMAGEM.")
            self._hide_format_row(); self.format_menu.config(values=[]); return
        self.current_mode = "image" if all_image else "video"
        values = self.image_formats if self.current_mode == "image" else self.video_formats
        if len(self.input_files) == 1:
            original_ext = _ext(self.input_files[0])
            if original_ext in values:
                values = [v for v in values if v != original_ext]
        self.format_menu.config(values=values)
        if self.formato_destino.get() not in values and values:
            self.formato_destino.set(values[0])
        if values: self._show_format_row()
        else: self._hide_format_row()

    # seleção/remover
    def selecionar_arquivos(self):
        tipos = [
            ("Vídeo/Áudio/Imagem", "*.mp4 *.avi *.mkv *.mov *.webm *.flv *.m4v *.wav *.mp3 *.jpg *.jpeg *.png *.webp *.bmp *.tif *.tiff *.cr2"),
            ("Vídeos", "*.mp4 *.avi *.mkv *.mov *.webm *.flv *.m4v"),
            ("Áudio", "*.wav *.mp3"),
            ("Imagens", "*.jpg *.jpeg *.png *.webp *.bmp *.tif *.tiff *.cr2"),
            ("Todos", "*.*"),
        ]
        paths = filedialog.askopenfilenames(title="Selecione arquivo(s)", filetypes=tipos)
        if paths: self._set_selected_files(list(paths))

    def _collect_supported_files(self, paths_or_dirs):
        collected = []
        for item in paths_or_dirs:
            if os.path.isfile(item):
                if _ext(item) in ALL_EXTS:
                    collected.append(os.path.abspath(item))
            elif os.path.isdir(item):
                try:
                    for name in os.listdir(item):
                        p = os.path.join(item, name)
                        if os.path.isfile(p) and _ext(p) in ALL_EXTS:
                            collected.append(os.path.abspath(p))
                except Exception:
                    pass
        seen = set(); unique = []
        for p in collected:
            if p.lower() not in seen:
                seen.add(p.lower()); unique.append(p)
        return unique

    def _on_drop_files(self, event):
        raw_items = self.tk.splitlist(event.data)
        if not raw_items: return
        paths = self._collect_supported_files(raw_items)
        if paths: self._set_selected_files(paths)

    def _set_selected_files(self, caminhos):
        self.input_files = list(caminhos)
        if not self.input_files:
            self.label_video.config(text="Nenhum arquivo selecionado")
            self.label_formato.config(text=""); self._update_format_menu(); return
        if len(self.input_files) == 1:
            f = self.input_files[0]
            self.label_video.config(text=f"Arquivo: {os.path.basename(f)}")
            self.label_formato.config(text=f"Formato original: {_ext(f)}")
        else:
            first = os.path.basename(self.input_files[0])
            self.label_video.config(text=f"{len(self.input_files)} arquivos selecionados (ex.: {first})")
            self.label_formato.config(text="Formatos originais variados")
        self._update_format_menu()

    def remover_arquivos(self):
        self.input_files = []
        self.label_video.config(text="Nenhum arquivo selecionado")
        self.label_formato.config(text="")
        self.progress_var.set(0); self.status_var.set("")
        self.open_btn.config(state=DISABLED)
        self.current_mode = "video"; self.formato_destino.set("mp4")
        self._hide_format_row(); self.format_menu.config(values=[])

    # conversão
    def start_conversion(self):
        if self.is_converting: return
        if not self.input_files:
            messagebox.showerror("Erro", "Selecione arquivo(s) primeiro."); return
        all_video = all(is_video_file(p) for p in self.input_files)
        all_image = all(is_image_file(p) for p in self.input_files)
        if not (all_video or all_image):
            messagebox.showerror("Seleção inválida", "Selecione apenas VÍDEO/ÁUDIO ou apenas IMAGEM."); return
        formato_destino = self.formato_destino.get()
        if not formato_destino:
            messagebox.showerror("Erro", "Selecione um formato de saída."); return

        self.is_converting = True; self.cancel_requested = False
        self._current_output_path = None
        self.convert_btn.config(state=DISABLED)
        self.cancel_btn.config(state=NORMAL)
        self.open_btn.config(state=DISABLED)
        self.progress_var.set(0); self.status_var.set("Preparando...")
        self.on_status("Conversão iniciada…")

        t = threading.Thread(target=self._batch_convert_worker, args=(self.input_files[:], formato_destino), daemon=True)
        t.start()

    def cancel_conversion(self):
        if self.is_converting:
            self.cancel_requested = True
            if self.proc and self.proc.poll() is None:
                try: self.proc.terminate()
                except Exception: pass

    def _batch_convert_worker(self, files, out_ext):
        try:
            total = len(files); last_out = None
            for idx, in_path in enumerate(files, start=1):
                if self.cancel_requested: break
                mode = "image" if is_image_file(in_path) and not is_video_file(in_path) else "video"
                pasta_saida = os.path.dirname(in_path)
                base = os.path.splitext(os.path.basename(in_path))[0]
                out_path = os.path.join(pasta_saida, base + f".{out_ext}")
                self._current_output_path = out_path
                self.ui_queue.put(("status", f"[{idx}/{total}] Preparando..."))
                ok = self._convert_single_image(in_path, out_path, idx, total) if mode=="image" else \
                     self._convert_single_video(in_path, out_path, idx, total)
                if not ok:
                    if self.cancel_requested: break
                    continue
                last_out = out_path
            if self.cancel_requested:
                self.ui_queue.put(("canceled", "Conversão cancelada."))
            else:
                if last_out: self.ultimo_arquivo_convertido = last_out
                self.ui_queue.put(("done", f"Conversão concluída de {total} arquivo(s)."))
        except Exception as e:
            if self.cancel_requested: self.ui_queue.put(("canceled","Conversão cancelada."))
            else: self.ui_queue.put(("error", f"Erro no processamento em lote: {e}"))
        finally:
            self.proc = None; self.is_converting = False

    def _convert_single_video(self, in_path, out_path, idx, total):
        try:
            duration = self._probe_duration(in_path)
            total_seconds = float(duration) if duration else None
            if out_path.lower().endswith(".mp3"):
                cmd = ["ffmpeg", "-y", "-i", in_path, "-vn", "-acodec", "libmp3lame", "-q:a", "2", out_path]
            else:
                cmd = ["ffmpeg", "-y", "-i", in_path, out_path]
            creationflags = create_no_window_flags()
            self.proc = subprocess.Popen(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True, creationflags=creationflags)
            for line in self.proc.stderr:
                if self.cancel_requested: break
                line = line.strip()
                if "time=" in line:
                    try:
                        t_str = line.split("time=")[1].split(" ")[0]
                        h,m,s = t_str.split(":")
                        sec = float(h)*3600 + float(m)*60 + float(s)
                        if total_seconds and total_seconds > 0:
                            pct = max(0.0, min(100.0, (sec/total_seconds)*100.0))
                            self.ui_queue.put(("progress", pct))
                            self.ui_queue.put(("status", f"[{idx}/{total}] Convertendo... {pct:.1f}% ({seconds_to_hms(sec)} de {seconds_to_hms(total_seconds)})"))
                        else:
                            self.ui_queue.put(("status", f"[{idx}/{total}] Convertendo... {seconds_to_hms(sec)}"))
                    except Exception:
                        pass
            ret = self.proc.wait()
            if self.cancel_requested:
                try:
                    if out_path and os.path.exists(out_path): os.remove(out_path)
                except Exception: pass
                return False
            if ret == 0:
                self.ui_queue.put(("status", f"[{idx}/{total}] Arquivo convertido: {os.path.basename(out_path)}"))
                return True
            else:
                self.ui_queue.put(("error", f"Falha na conversão do arquivo: {os.path.basename(in_path)}"))
                return False
        except FileNotFoundError:
            self.ui_queue.put(("error", "Não encontrei o ffmpeg/ffprobe. Instale-os e adicione ao PATH."))
            return False
        except Exception as e:
            if self.cancel_requested:
                try:
                    if out_path and os.path.exists(out_path): os.remove(out_path)
                except Exception: pass
                return False
            else:
                self.ui_queue.put(("error", f"Erro: {e}"))
                return False
        finally:
            self.proc = None

    def _convert_single_image(self, in_path, out_path, idx, total):
        try:
            if not HAS_PIL:
                self.ui_queue.put(("error", "Conversão de imagens requer Pillow. Instale: pip install pillow"))
                return False
            self.ui_queue.put(("status", f"[{idx}/{total}] Convertendo imagem..."))
            src_ext = _ext(in_path); dst_ext = _ext(out_path)
            if src_ext == "cr2":
                if not HAS_RAWPY:
                    self.ui_queue.put(("error", "Para converter CR2 instale rawpy: pip install rawpy"))
                    return False
                with rawpy.imread(in_path) as raw:
                    rgb = raw.postprocess()
                im = Image.fromarray(rgb)
            else:
                im = Image.open(in_path)
            save_kwargs = {}; fmt = dst_ext.upper()
            if dst_ext in ("jpg","jpeg"):
                if im.mode in ("RGBA","LA","P"): im = im.convert("RGB")
                save_kwargs["quality"] = 92; fmt = "JPEG"
            elif dst_ext == "webp":
                fmt = "WEBP"; save_kwargs["quality"] = 92
            elif dst_ext in ("tif","tiff"): fmt = "TIFF"
            elif dst_ext == "bmp": fmt = "BMP"
            im.save(out_path, fmt, **save_kwargs)
            if self.cancel_requested:
                try:
                    if out_path and os.path.exists(out_path): os.remove(out_path)
                except Exception: pass
                return False
            self.ui_queue.put(("progress", 100))
            self.ui_queue.put(("status", f"[{idx}/{total}] Arquivo convertido: {os.path.basename(out_path)}"))
            return True
        except Exception as e:
            if self.cancel_requested:
                try:
                    if out_path and os.path.exists(out_path): os.remove(out_path)
                except Exception: pass
                return False
            else:
                self.ui_queue.put(("error", f"Erro na conversão de imagem ({os.path.basename(in_path)}): {e}"))
                return False

    def _probe_duration(self, path):
        try:
            creationflags = create_no_window_flags()
            cmd = ["ffprobe", "-v", "error", "-show_entries", "format=duration", "-of", "default=nk=1:nw=1", path]
            out = subprocess.check_output(cmd, text=True, creationflags=creationflags, stderr=subprocess.DEVNULL)
            return out.strip()
        except Exception:
            return None

    def _drain_ui_queue(self):
        try:
            while True:
                kind, payload = self.ui_queue.get_nowait()
                if kind == "progress":
                    self.progress_var.set(payload)
                elif kind == "status":
                    self.status_var.set(payload)
                    self.on_status(payload)
                elif kind == "done":
                    self.is_converting = False
                    self.progress_var.set(100)
                    self.status_var.set(payload)
                    self.on_status("Conversão finalizada")
                    self.convert_btn.config(state=NORMAL)
                    self.cancel_btn.config(state=DISABLED)
                    self.open_btn.config(state=NORMAL)
                    messagebox.showinfo("Sucesso", "Conversão concluída!")
                elif kind == "canceled":
                    self.is_converting = False
                    self.status_var.set(payload)
                    self.on_status(payload)
                    self.progress_var.set(0)
                    self.convert_btn.config(state=NORMAL)
                    self.cancel_btn.config(state=DISABLED)
                    self.open_btn.config(state=DISABLED)
                    self._current_output_path = None
                elif kind == "error":
                    self.on_status("Erro no conversor")
                    messagebox.showerror("Erro", str(payload))
        except queue.Empty:
            pass
        finally:
            if not self.is_converting and self.cancel_btn["state"] == NORMAL:
                self.convert_btn.config(state=NORMAL)
                self.cancel_btn.config(state=DISABLED)
            self.after(100, self._drain_ui_queue)

    def abrir_pasta(self):
        if self.ultimo_arquivo_convertido and os.path.exists(self.ultimo_arquivo_convertido):
            pasta = os.path.dirname(self.ultimo_arquivo_convertido)
            try:
                if sys.platform.startswith("win"): os.startfile(pasta)
                elif sys.platform == "darwin": subprocess.Popen(["open", pasta])
                else: subprocess.Popen(["xdg-open", pasta])
            except Exception as e:
                messagebox.showerror("Erro", f"Não foi possível abrir a pasta: {e}")
        else:
            messagebox.showerror("Erro", "Nenhum arquivo convertido encontrado.")

# ======================= Transcrição =======================
AUDIO_VIDEO_EXTS = {"mp3", "wav", "m4a", "mp4", "mkv", "mov", "webm"}
SUPPORTED_EXTS = AUDIO_VIDEO_EXTS
WHISPER_MODEL = "small"  # ajuste se quiser

class TranscriberFrame(ttk.Frame):
    def __init__(self, master, on_status):
        super().__init__(master)
        self.on_status = on_status
        self.input_files = []
        self.progress_var = tk.DoubleVar(value=0)
        self.status_var = tk.StringVar(value="")
        self.is_running = False
        self.cancel_requested = False
        self.ui_queue = queue.Queue()
        self.last_output = None
        self.model = None

        self._build_ui()
        self.after(100, self._drain_ui_queue)

        if HAS_DND:
            try:
                self.drop_target_register(DND_FILES)
                self.dnd_bind("<<Drop>>", self._on_drop_files)
            except Exception:
                pass

    def _build_ui(self):
        card = ttk.Frame(self, padding=18, bootstyle="dark")
        card.pack(fill="both", expand=True)

        header = ttk.Frame(card)
        header.pack(fill="x")
        ttk.Label(header, text="📝  Transcritor de Áudio (Word)", font=("Helvetica", 18, "bold")).pack(side="left")
        ttk.Separator(card).pack(fill="x", pady=12)

        row = ttk.Frame(card)
        row.pack(fill="x")
        ttk.Button(row, text="Selecionar Arquivo(s)", command=self.selecionar_arquivos, bootstyle=WARNING).pack(side="left")
        ttk.Button(row, text="🗑 Remover", command=self.remover_arquivos, bootstyle=DANGER).pack(side="left", padx=(10, 0))

        self.label_sel = ttk.Label(card, text="Nenhum arquivo selecionado", font=("Helvetica", 12))
        self.label_sel.pack(anchor="w", pady=(10, 0))

        ctl = ttk.Frame(card)
        ctl.pack(fill="x", pady=(10, 6))
        self.btn_run = ttk.Button(ctl, text="▶️ Transcrever", command=self.start_transcription, bootstyle=SUCCESS)
        self.btn_run.pack(side="left")
        self.btn_cancel = ttk.Button(ctl, text="Cancelar", command=self.cancel_transcription, bootstyle=SECONDARY, state=DISABLED)
        self.btn_cancel.pack(side="left", padx=(10, 0))

        prog_card = ttk.Frame(card, padding=10, bootstyle="secondary")
        prog_card.pack(fill="x", pady=(8, 4))
        self.progress = ttk.Progressbar(prog_card, orient=tk.HORIZONTAL, mode="determinate", variable=self.progress_var, maximum=100)
        self.progress.pack(fill="x")
        self.status_lbl = ttk.Label(prog_card, textvariable=self.status_var, font=("Helvetica", 11))
        self.status_lbl.pack(anchor="w", pady=(6,0))

        self.btn_open = ttk.Button(card, text="Abrir pasta do último .docx", command=self.abrir_pasta, bootstyle=INFO, state=DISABLED)
        self.btn_open.pack(pady=8)

    def selecionar_arquivos(self):
        tipos = [
            ("Áudio/Vídeo", "*.mp3 *.wav *.m4a *.mp4 *.mkv *.mov *.webm"),
            ("Áudio", "*.mp3 *.wav *.m4a"),
            ("Vídeo", "*.mp4 *.mkv *.mov *.webm"),
            ("Todos", "*.*"),
        ]
        paths = filedialog.askopenfilenames(title="Selecione arquivo(s)", filetypes=tipos)
        if paths: self._set_files(list(paths))

    def _on_drop_files(self, event):
        items = self.tk.splitlist(event.data)
        paths = []
        for p in items:
            if os.path.isdir(p):
                try:
                    for nm in os.listdir(p):
                        f = os.path.join(p, nm)
                        if os.path.isfile(f) and _ext(f) in SUPPORTED_EXTS:
                            paths.append(os.path.abspath(f))
                except Exception:
                    pass
            elif os.path.isfile(p) and _ext(p) in SUPPORTED_EXTS:
                paths.append(os.path.abspath(p))
        seen, uniq = set(), []
        for p in paths:
            low = p.lower()
            if low not in seen:
                seen.add(low); uniq.append(p)
        if uniq: self._set_files(uniq)

    def _set_files(self, paths):
        paths = [p for p in paths if os.path.isfile(p) and _ext(p) in SUPPORTED_EXTS]
        self.input_files = paths
        if not self.input_files:
            self.label_sel.config(text="Nenhum arquivo selecionado"); return
        if len(paths) == 1:
            self.label_sel.config(text=f"Arquivo: {os.path.basename(paths[0])}")
        else:
            self.label_sel.config(text=f"{len(paths)} arquivos selecionados (ex.: {os.path.basename(paths[0])})")

    def remover_arquivos(self):
        self.input_files = []
        self.label_sel.config(text="Nenhum arquivo selecionado")
        self.progress_var.set(0); self.status_var.set("")
        self.btn_open.config(state=DISABLED)

    def start_transcription(self):
        if self.is_running: return
        if _MISSING_TRANSCRIBE_DEPS:
            messagebox.showerror("Dependências", f"Instale: pip install {' '.join(_MISSING_TRANSCRIBE_DEPS)}")
            return
        if not self.input_files:
            messagebox.showerror("Erro", "Selecione pelo menos um arquivo de áudio/vídeo."); return

        self.is_running = True; self.cancel_requested = False
        self.btn_run.config(state=DISABLED); self.btn_cancel.config(state=NORMAL)
        self.btn_open.config(state=DISABLED)
        self.progress_var.set(0); self.status_var.set("Preparando...")
        self.on_status("Transcrição iniciada…")

        t = threading.Thread(target=self._batch_transcribe_worker, args=(self.input_files[:],), daemon=True)
        t.start()

    def cancel_transcription(self):
        if self.is_running: self.cancel_requested = True

    def _ensure_model(self):
        if self.model is not None: return
        self.ui_queue.put(("status", f"Carregando modelo '{WHISPER_MODEL}' na CPU (int8)..."))
        self.model = WhisperModel(WHISPER_MODEL, device="cpu", compute_type="int8")
        self.ui_queue.put(("status", "Modelo carregado na CPU."))

    def _batch_transcribe_worker(self, files):
        try:
            self._ensure_model()
            total = len(files); last_out = None
            for idx, path in enumerate(files, start=1):
                if self.cancel_requested: break
                base = os.path.splitext(os.path.basename(path))[0]
                out_path = os.path.join(os.path.dirname(path), base + ".docx")
                self.ui_queue.put(("status", f"[{idx}/{total}] Transcrevendo: {os.path.basename(path)}"))
                ok = self._transcribe_one(path, out_path, idx, total)
                if ok: last_out = out_path
            if self.cancel_requested:
                self.ui_queue.put(("canceled", "Transcrição cancelada."))
            else:
                if last_out: self.last_output = last_out
                self.ui_queue.put(("done", f"Transcrição concluída de {total} arquivo(s)."))
        except Exception as e:
            if self.cancel_requested: self.ui_queue.put(("canceled","Transcrição cancelada."))
            else: self.ui_queue.put(("error", f"Erro no processamento: {e}"))
        finally:
            self.is_running = False

    def _transcribe_one(self, in_path, out_docx, idx, total):
        try:
            segments, info = self.model.transcribe(
                in_path,
                language=None,  # autodetect
                vad_filter=True,
                beam_size=5,
                condition_on_previous_text=True
            )
            pieces = []
            for seg in segments:
                if self.cancel_requested: break
                txt = (seg.text or "").strip()
                if txt: pieces.append(txt)
            if self.cancel_requested: return False
            full_text = " ".join(pieces).strip()
            if not full_text:
                self.ui_queue.put(("error", f"Nenhum texto reconhecido em: {os.path.basename(in_path)}"))
                return False
            try:
                doc = Document()
                doc.add_heading("Transcrição", level=1)
                if getattr(info, "language", None):
                    lang = info.language
                    doc.add_paragraph(f"Idioma detectado: {lang}")
                    doc.add_paragraph("")
                doc.add_paragraph(full_text)
                doc.save(out_docx)
            except Exception as e:
                self.ui_queue.put(("error", f"Falha ao salvar DOCX ({os.path.basename(in_path)}): {e}"))
                return False
            self.ui_queue.put(("progress", 100))
            self.ui_queue.put(("status", f"[{idx}/{total}] Salvo: {os.path.basename(out_docx)}"))
            return True
        except Exception as e:
            self.ui_queue.put(("error", f"Erro ao transcrever {os.path.basename(in_path)}: {e}"))
            return False

    def _drain_ui_queue(self):
        try:
            while True:
                kind, payload = self.ui_queue.get_nowait()
                if kind == "progress":
                    self.progress_var.set(payload)
                elif kind == "status":
                    self.status_var.set(payload)
                    self.on_status(payload)
                elif kind == "done":
                    self.progress_var.set(100)
                    self.status_var.set(payload)
                    self.on_status("Transcrição finalizada")
                    self.btn_run.config(state=NORMAL)
                    self.btn_cancel.config(state=DISABLED)
                    self.btn_open.config(state=NORMAL if self.last_output else DISABLED)
                    messagebox.showinfo("Sucesso", "Transcrição concluída!")
                elif kind == "canceled":
                    self.status_var.set(payload)
                    self.on_status(payload)
                    self.progress_var.set(0)
                    self.btn_run.config(state=NORMAL)
                    self.btn_cancel.config(state=DISABLED)
                    self.btn_open.config(state=DISABLED)
                elif kind == "error":
                    self.on_status("Erro na transcrição")
                    messagebox.showerror("Erro", str(payload))
        except queue.Empty:
            pass
        finally:
            if not self.is_running and self.btn_cancel["state"] == NORMAL:
                self.btn_run.config(state=NORMAL)
                self.btn_cancel.config(state=DISABLED)
            self.after(100, self._drain_ui_queue)

    def abrir_pasta(self):
        if self.last_output and os.path.exists(self.last_output):
            pasta = os.path.dirname(self.last_output)
            try:
                if sys.platform.startswith("win"): os.startfile(pasta)
                elif sys.platform == "darwin": subprocess.Popen(["open", pasta])
                else: subprocess.Popen(["xdg-open", pasta])
            except Exception as e:
                messagebox.showerror("Erro", f"Não foi possível abrir a pasta: {e}")
        else:
            messagebox.showerror("Erro", "Nenhum arquivo transcrito encontrado.")

# ======================= YouTube =======================
CONFIG_FILE = Path("config.json")

def ensure_yt_dlp_updated():
    try:
        subprocess.run([sys.executable, "-m", "pip", "install", "-U", "yt-dlp"], check=False, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
    except Exception:
        pass

ensure_yt_dlp_updated()
import yt_dlp
from yt_dlp.utils import DownloadCancelled

class YouTubeFrame(ttk.Frame):
    def __init__(self, master, on_status):
        super().__init__(master)
        self.on_status = on_status
        self.destination_folder = self.load_config()
        self.selected_format = tk.StringVar(value="Música")
        self.selected_quality = tk.StringVar(value="best")
        self.downloaded_file = None
        self.cancel_requested = False
        self._last_tmp_file = None

        self._build_ui()
        self._apply_quality_visibility()

    def _build_ui(self):
        card = ttk.Frame(self, padding=18, bootstyle="dark")
        card.pack(fill="both", expand=True)

        header = ttk.Frame(card)
        header.pack(fill="x")
        ttk.Label(header, text="⬇️  Baixar do YouTube", font=("Helvetica", 18, "bold")).pack(side="left")
        ttk.Separator(card).pack(fill="x", pady=12)

        url_row = ttk.Frame(card)
        url_row.pack(fill="x")
        ttk.Label(url_row, text="YouTube URL:", font=("Helvetica", 13)).pack(side="left")
        self.url_entry = ttk.Entry(url_row, width=60, font=("Helvetica", 12))
        self.url_entry.pack(side="left", fill="x", expand=True, padx=(10,0))

        dest_row = ttk.Frame(card)
        dest_row.pack(fill="x", pady=(10,0))
        ttk.Button(dest_row, text="Escolher pasta de destino", command=self.choose_dest_folder, bootstyle=SUCCESS).pack(side="left")
        self.dest_label = ttk.Label(dest_row, text=self.destination_folder or "Nenhuma pasta selecionada", anchor="w", font=("Helvetica", 12))
        self.dest_label.pack(side="left", fill="x", expand=True, padx=(10,0))

        opts = ttk.Frame(card)
        opts.pack(fill="x", pady=(10,0))
        ttk.Label(opts, text="Formato:", font=("Helvetica", 13)).pack(side="left")
        self.format_menu = ttk.Combobox(opts, textvariable=self.selected_format, values=["Música", "Vídeo"], state="readonly", width=12)
        self.format_menu.pack(side="left", padx=(8, 20))
        self.format_menu.bind("<<ComboboxSelected>>", self._on_format_change)

        self.quality_label = ttk.Label(opts, text="Qualidade do vídeo:", font=("Helvetica", 13))
        self.quality_label.pack(side="left")
        self.quality_menu = ttk.Combobox(opts, textvariable=self.selected_quality, values=["best", "144p", "240p", "360p", "480p","720p","1080p","1440p","2160p"], state="readonly", width=12)
        self.quality_menu.pack(side="left", padx=(8, 0))
        self._quality_widgets = [self.quality_label, self.quality_menu]

        ctl = ttk.Frame(card); ctl.pack(fill="x", pady=(14, 8))
        self.download_btn = ttk.Button(ctl, text="▶️ Baixar", command=self.start_download, bootstyle=PRIMARY)
        self.download_btn.pack(side="left")
        self.cancel_btn = ttk.Button(ctl, text="Cancelar", command=self.cancel_download, bootstyle=SECONDARY, state=DISABLED)
        self.cancel_btn.pack(side="left", padx=(10, 0))

        prog_card = ttk.Frame(card, padding=10, bootstyle="secondary")
        prog_card.pack(fill="x", pady=(8, 4))
        self.progress = ttk.Progressbar(prog_card, orient=tk.HORIZONTAL, mode="determinate")
        self.progress.pack(fill="x")
        self.status = ttk.Label(prog_card, text="", font=("Helvetica", 11))
        self.status.pack(anchor="w", pady=(6,0))

        self.open_folder_button = ttk.Button(card, text="Abrir local do arquivo", command=self.open_file_location, bootstyle=INFO, state=DISABLED)
        self.open_folder_button.pack(pady=8)

    # config
    def load_config(self):
        if CONFIG_FILE.exists():
            with CONFIG_FILE.open("r", encoding="utf-8") as f:
                return json.load(f).get("destination_folder", "")
        return ""

    def save_config(self):
        with CONFIG_FILE.open("w", encoding="utf-8") as f:
            json.dump({"destination_folder": self.destination_folder}, f)

    def choose_dest_folder(self):
        folder = filedialog.askdirectory()
        if folder:
            self.destination_folder = folder
            self.dest_label.config(text=folder)
            self.save_config()

    # eventos UI
    def _on_format_change(self, _evt=None):
        self._apply_quality_visibility()

    def _is_video(self): return self.selected_format.get() == "Vídeo"

    def _apply_quality_visibility(self):
        for w in self._quality_widgets:
            try:
                if self._is_video(): w.pack()
                else: w.pack_forget()
            except Exception:
                pass

    # download
    def start_download(self):
        url = self.url_entry.get().strip()
        if not url:
            messagebox.showerror("Erro", "Insira a URL do YouTube."); return
        if not self.destination_folder:
            messagebox.showerror("Erro", "Escolha a pasta de destino."); return

        self.progress["value"] = 0
        self.status.config(text="Preparando...")
        self.open_folder_button.config(state=DISABLED)
        self.downloaded_file = None
        self.cancel_requested = False
        self._last_tmp_file = None
        self.download_btn.config(state=DISABLED)
        self.cancel_btn.config(state=NORMAL)
        self.on_status("Download iniciado…")

        fmt_mode = self.selected_format.get()
        qual = self.selected_quality.get()
        threading.Thread(target=self.download_media, args=(url, fmt_mode, qual), daemon=True).start()

    def cancel_download(self):
        if self.cancel_requested: return
        self.cancel_requested = True
        self.cancel_btn.config(state=DISABLED)
        self.status.config(text="Cancelando...")
        self.on_status("Cancelando download…")

    def _cleanup_partial(self):
        paths = set()
        if self._last_tmp_file:
            paths.add(self._last_tmp_file)
            if not self._last_tmp_file.endswith(".part"):
                paths.add(self._last_tmp_file + ".part")
            paths.add(self._last_tmp_file + ".ytdl")
        for p in list(paths):
            try:
                if p and os.path.exists(p): os.remove(p)
            except Exception:
                pass

    def download_media(self, url, fmt_mode, quality_choice):
        try:
            # Workaround para SABR: força clientes sem web/web_safari
            common_args = {
                "extractor_args": {
                    "youtube": {
                        "player_client": ["android"],           # alternativas: "tv_embedded", "ios"
                        "player_skip": ["web_safari", "web"]
                    }
                },
                "noprogress": True, "nocolor": True, "quiet": True,
                "progress_hooks": [self.ydl_hook],
                "outtmpl": os.path.join(self.destination_folder, "%(title)s.%(ext)s"),
            }

            if fmt_mode == "Música":
                ydl_opts = {
                    **common_args,
                    "format": "bestaudio/best",
                    "postprocessors": [{"key": "FFmpegExtractAudio", "preferredcodec": "mp3"}],
                }
            else:
                if quality_choice == "best":
                    # melhor disponível em MP4
                    vsel = "bestvideo[ext=mp4]"
                    fmt = f"{vsel}+bestaudio[ext=m4a]/best[ext=mp4]/best"
                    ydl_opts = {**common_args, "format": fmt, "merge_output_format": "mp4"}
                else:
                    # força a priorização da resolução escolhida (ou a mais próxima abaixo)
                    h = "".join(ch for ch in quality_choice if ch.isdigit()) or "1080"
                    vsel = f"bestvideo[ext=mp4][height<={h}]"
                    fmt = f"{vsel}+bestaudio[ext=m4a]/best[ext=mp4]/best"
                    ydl_opts = {
                        **common_args,
                        "format": fmt,
                        "merge_output_format": "mp4",
                        # ordena forçando preferência pela resolução alvo
                        "format_sort": [f"res:{h}"],
                        "format_sort_force": True,
                    }

            with yt_dlp.YoutubeDL(ydl_opts) as ydl:
                try:
                    ydl.download([url])
                except DownloadCancelled:
                    self._cleanup_partial(); self._finish_canceled(); return
            self._finish_ok()
        except Exception as e:
            if self.cancel_requested:
                self._cleanup_partial(); self._finish_canceled(); return
            self._finish_error(str(e))

    def ydl_hook(self, d):
        if self.cancel_requested: raise DownloadCancelled()
        status = d.get("status")
        self._last_tmp_file = d.get("tmpfilename") or d.get("filename") or self._last_tmp_file
        if status == "finished":
            self.progress["value"] = 100
            self.status.config(text="Concluído!")
            self.downloaded_file = d.get("filename") or d.get("info_dict", {}).get("_filename")
            self.on_status("Download finalizado")
            return
        if status == "downloading":
            downloaded = d.get("downloaded_bytes")
            total = d.get("total_bytes") or d.get("total_bytes_estimate")
            speed = d.get("speed")
            pct = None
            if downloaded is not None and total:
                try: pct = max(0.0, min(100.0, (downloaded/total)*100.0))
                except Exception: pct = None
            if pct is not None: self.progress["value"] = pct
            parts = ["Baixando..."]
            if pct is not None: parts.append(f"{pct:.1f}%")
            if downloaded is not None:
                if total: parts.append(f"({format_bytes(downloaded)} de {format_bytes(total)})")
                else: parts.append(f"({format_bytes(downloaded)})")
            if speed:
                sp = format_bytes(speed)
                if sp: parts.append(f"{sp}/s")
            msg = " • ".join([p for p in parts if p])
            self.status.config(text=msg)
            self.on_status(msg)

    def _finish_ok(self):
        self.status.config(text="Download concluído!")
        self.open_folder_button.config(state=NORMAL)
        self.download_btn.config(state=NORMAL)
        self.cancel_btn.config(state=DISABLED)

    def _finish_canceled(self):
        self.status.config(text="Download cancelado.")
        self.progress["value"] = 0
        self.open_folder_button.config(state=DISABLED)
        self.download_btn.config(state=NORMAL)
        self.cancel_btn.config(state=DISABLED)
        self.on_status("Download cancelado")

    def _finish_error(self, msg):
        self.status.config(text="")
        self.download_btn.config(state=NORMAL)
        self.cancel_btn.config(state=DISABLED)
        self.on_status("Erro no download")
        messagebox.showerror("Erro", msg)

    def open_file_location(self):
        if self.downloaded_file:
            file_dir = os.path.dirname(self.downloaded_file)
            try:
                if sys.platform.startswith("win"): os.startfile(file_dir)
                elif sys.platform == "darwin": subprocess.Popen(["open", file_dir])
                else: subprocess.Popen(["xdg-open", file_dir])
            except Exception as e:
                messagebox.showerror("Erro", f"Não foi possível abrir o local do arquivo: {e}")
        else:
            messagebox.showerror("Erro", "Nenhum arquivo baixado encontrado.")

# =================== App principal com seletor de módulos ===================
class SuperApp(ttk.Window if not HAS_DND else TkinterDnD.Tk):
    def __init__(self):
        if HAS_DND:
            super().__init__()
            self.title("Mídia Suite — Conversor • YouTube • Transcrição")
            self.style = ttk.Style(theme="darkly")
            self.geometry("1040x620")
        else:
            super().__init__(title="Mídia Suite — Conversor • YouTube • Transcrição", themename="darkly", size=(1040, 620))
        self.minsize(980, 580)

        # Layout raiz
        self.grid_columnconfigure(1, weight=1)
        self.grid_rowconfigure(1, weight=1)

        # Header topbar
        topbar = ttk.Frame(self, padding=(16, 12))
        topbar.grid(row=0, column=0, columnspan=2, sticky="ew")
        ttk.Label(topbar, text="Mídia Suite", font=("Helvetica", 22, "bold")).pack(side="left")
        self.title_label = ttk.Label(topbar, text=" — Conversor", font=("Helvetica", 16))
        self.title_label.pack(side="left")

        # Sidebar
        sidebar = ttk.Frame(self, padding=16)
        sidebar.grid(row=1, column=0, sticky="ns")
        ttk.Button(sidebar, text="⚙️  Conversor", bootstyle=PRIMARY, command=lambda: self._show("converter")).pack(pady=6, fill="x")
        ttk.Button(sidebar, text="⬇️  YouTube", bootstyle=INFO, command=lambda: self._show("youtube")).pack(pady=6, fill="x")
        ttk.Button(sidebar, text="📝  Transcrição", bootstyle=SUCCESS, command=lambda: self._show("transcribe")).pack(pady=6, fill="x")

        # Área de conteúdo (cartão grande)
        self.content = ttk.Frame(self, padding=(6, 16, 16, 16))
        self.content.grid(row=1, column=1, sticky="nsew")
        self.content.grid_columnconfigure(0, weight=1)
        self.content.grid_rowconfigure(0, weight=1)

        # Statusbar bottom
        sb = ttk.Frame(self, padding=(16, 8))
        sb.grid(row=2, column=0, columnspan=2, sticky="ew")
        self.statusbar_var = tk.StringVar(value="Pronto.")
        ttk.Label(sb, textvariable=self.statusbar_var, anchor="w").pack(side="left")

        # Stacked frames
        self.frames = {
            "converter": ConverterFrame(self.content, self._set_status),
            "youtube": YouTubeFrame(self.content, self._set_status),
            "transcribe": TranscriberFrame(self.content, self._set_status),
        }
        for f in self.frames.values():
            f.grid(row=0, column=0, sticky="nsew")

        self._show("converter")

        # Atalhos úteis
        self.bind("<Control-Key-1>", lambda e: self._show("converter"))
        self.bind("<Control-Key-2>", lambda e: self._show("youtube"))
        self.bind("<Control-Key-3>", lambda e: self._show("transcribe"))

    def _show(self, key):
        f = self.frames.get(key)
        if not f: return
        f.lift()
        if key == "converter": self.title_label.config(text=" — Conversor")
        elif key == "youtube": self.title_label.config(text=" — YouTube")
        else: self.title_label.config(text=" — Transcrição")
        # feedback de seleção
        self._set_status("Pronto.")

    def _set_status(self, text):
        self.statusbar_var.set(text)

if __name__ == "__main__":
    app = SuperApp()
    app.mainloop()
