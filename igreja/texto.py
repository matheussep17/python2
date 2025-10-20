import threading
import tkinter as tk
from tkinter import filedialog, messagebox
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import os
import subprocess
import sys
import queue

# --- Drag and Drop (opcional) ---
try:
    from tkinterdnd2 import DND_FILES, TkinterDnD
    HAS_DND = True
except Exception:
    HAS_DND = False

# --- Dependências para transcrição/Word ---
_MISSING_DEPS = []
try:
    from faster_whisper import WhisperModel
    HAS_FW = True
except Exception:
    HAS_FW = False
    _MISSING_DEPS.append("faster-whisper")

try:
    from docx import Document
    HAS_DOCX = True
except Exception:
    HAS_DOCX = False
    _MISSING_DEPS.append("python-docx")


def create_no_window_flags():
    if sys.platform.startswith("win"):
        return subprocess.CREATE_NO_WINDOW
    return 0


def seconds_to_hms(s):
    try:
        s = float(s)
    except Exception:
        return "00:00:00"
    h = int(s // 3600)
    m = int((s % 3600) // 60)
    sec = int(s % 60)
    return f"{h:02d}:{m:02d}:{sec:02d}"


def _ext(path):
    return os.path.splitext(path)[1][1:].lower()


AUDIO_VIDEO_EXTS = {"mp3", "wav", "m4a", "mp4", "mkv", "mov", "webm"}
SUPPORTED_EXTS = AUDIO_VIDEO_EXTS

# Modelo (ajuste se quiser outro: "medium", "large-v3" etc.)
WHISPER_MODEL = "small"


class AudioTranscriberApp(ttk.Window if not HAS_DND else TkinterDnD.Tk):
    def __init__(self):
        if HAS_DND:
            super().__init__()
            self.title("Transcritor de Áudio (Word)")
            self.style = ttk.Style(theme="darkly")
            self.geometry("600x480")
        else:
            super().__init__(title="Transcritor de Áudio (Word)", themename="darkly", size=(600, 480))

        self.minsize(600, 420)
        self.center_window(600, 480)

        # Estado
        self.input_files = []
        self.progress_var = tk.DoubleVar(value=0)
        self.status_var = tk.StringVar(value="")
        self.is_running = False
        self.cancel_requested = False
        self.ui_queue = queue.Queue()
        self.last_output = None
        self.model = None  # cache do Whisper

        self.init_ui()
        self.after(100, self._drain_ui_queue)

        if HAS_DND:
            try:
                self.drop_target_register(DND_FILES)
                self.dnd_bind("<<Drop>>", self._on_drop_files)
            except Exception:
                pass

    # ----------- UI -----------
    def center_window(self, width, height):
        self.update_idletasks()
        sw = self.winfo_screenwidth()
        sh = self.winfo_screenheight()
        x = (sw - width) // 2
        y = (sh - height) // 2
        self.geometry(f"{width}x{height}+{x}+{y}")

    def init_ui(self):
        container = ttk.Frame(self, padding=20)
        container.pack(fill="both", expand=True)

        ttk.Label(container, text="Transcritor de Áudio (Word)", font=("Helvetica", 20, "bold")).pack(pady=(0, 10))

        row = ttk.Frame(container)
        row.pack(pady=5, fill="x")

        ttk.Button(row, text="Selecionar Arquivo(s)", command=self.selecionar_arquivos, bootstyle=WARNING).pack(side="left")
        ttk.Button(row, text="🗑 Remover arquivo(s)", command=self.remover_arquivos, bootstyle=DANGER).pack(side="left", padx=(10, 0))

        self.label_sel = ttk.Label(container, text="Nenhum arquivo selecionado", font=("Helvetica", 12))
        self.label_sel.pack(anchor="w", pady=(8, 0))

        if HAS_DND:
            ttk.Label(
                container,
                font=("Helvetica", 10, "italic"),
                foreground="#9aa0a6",
            ).pack(anchor="w", pady=(0, 6))

        # Botões Transcrever / Cancelar
        btns = ttk.Frame(container)
        btns.pack(pady=(10, 6), fill="x")
        self.btn_run = ttk.Button(btns, text="Transcrever", command=self.start_transcription, bootstyle=SUCCESS)
        self.btn_run.pack(side="left")
        self.btn_cancel = ttk.Button(btns, text="Cancelar", command=self.cancel_transcription, bootstyle=SECONDARY, state=DISABLED)
        self.btn_cancel.pack(side="left", padx=(10, 0))

        # Progresso + status
        self.progress = ttk.Progressbar(container, orient=tk.HORIZONTAL, length=400, mode="determinate",
                                        variable=self.progress_var, maximum=100)
        self.progress.pack(fill="x", pady=(6, 2))
        self.status_lbl = ttk.Label(container, textvariable=self.status_var, font=("Helvetica", 11))
        self.status_lbl.pack(anchor="w")

        # Abrir pasta
        self.btn_open = ttk.Button(container, text="Abrir pasta do último .docx", command=self.abrir_pasta, bootstyle=INFO, state=DISABLED)
        self.btn_open.pack(pady=10)

    # ----------- Seleção / DnD -----------
    def selecionar_arquivos(self):
        tipos = [
            ("Áudio/Vídeo", "*.mp3 *.wav *.m4a *.mp4 *.mkv *.mov *.webm"),
            ("Áudio", "*.mp3 *.wav *.m4a"),
            ("Vídeo", "*.mp4 *.mkv *.mov *.webm"),
            ("Todos", "*.*"),
        ]
        paths = filedialog.askopenfilenames(title="Selecione arquivo(s)", filetypes=tipos)
        if paths:
            self._set_files(list(paths))

    def _on_drop_files(self, event):
        items = self.tk.splitlist(event.data)
        paths = []
        for p in items:
            if os.path.isdir(p):
                # pega arquivos do nível da pasta (sem recursão)
                try:
                    for nm in os.listdir(p):
                        f = os.path.join(p, nm)
                        if os.path.isfile(f) and _ext(f) in SUPPORTED_EXTS:
                            paths.append(os.path.abspath(f))
                except Exception:
                    pass
            elif os.path.isfile(p) and _ext(p) in SUPPORTED_EXTS:
                paths.append(os.path.abspath(p))
        # dedup preservando ordem
        seen, uniq = set(), []
        for p in paths:
            low = p.lower()
            if low not in seen:
                seen.add(low)
                uniq.append(p)
        if uniq:
            self._set_files(uniq)

    def _set_files(self, paths):
        paths = [p for p in paths if os.path.isfile(p) and _ext(p) in SUPPORTED_EXTS]
        self.input_files = paths
        if not self.input_files:
            self.label_sel.config(text="Nenhum arquivo selecionado")
            return
        if len(paths) == 1:
            self.label_sel.config(text=f"Arquivo: {os.path.basename(paths[0])}")
        else:
            self.label_sel.config(text=f"{len(paths)} arquivos selecionados (ex.: {os.path.basename(paths[0])})")

    def remover_arquivos(self):
        self.input_files = []
        self.label_sel.config(text="Nenhum arquivo selecionado")
        self.progress_var.set(0)
        self.status_var.set("")
        self.btn_open.config(state=DISABLED)

    # ----------- Execução -----------
    def start_transcription(self):
        if self.is_running:
            return
        if _MISSING_DEPS:
            messagebox.showerror("Dependências", f"Instale: pip install {' '.join(_MISSING_DEPS)}")
            return
        if not self.input_files:
            messagebox.showerror("Erro", "Selecione pelo menos um arquivo de áudio/vídeo.")
            return

        self.is_running = True
        self.cancel_requested = False
        self.btn_run.config(state=DISABLED)
        self.btn_cancel.config(state=NORMAL)
        self.btn_open.config(state=DISABLED)
        self.progress_var.set(0)
        self.status_var.set("Preparando...")

        t = threading.Thread(target=self._batch_transcribe_worker, args=(self.input_files[:],), daemon=True)
        t.start()

    def cancel_transcription(self):
        if self.is_running:
            self.cancel_requested = True

    # Carrega modelo FORÇANDO CPU (sem CUDA/cuDNN)
    def _ensure_model(self):
        if self.model is not None:
            return
        self.ui_queue.put(("status", f"Carregando modelo '{WHISPER_MODEL}' na CPU (int8)..."))
        # Força CPU aqui: evita qualquer tentativa de usar GPU e cuDNN
        self.model = WhisperModel(WHISPER_MODEL, device="cpu", compute_type="int8")
        self.ui_queue.put(("status", "Modelo carregado na CPU."))

    # Worker em lote
    def _batch_transcribe_worker(self, files):
        try:
            self._ensure_model()
            total = len(files)
            last_out = None

            for idx, path in enumerate(files, start=1):
                if self.cancel_requested:
                    break

                base = os.path.splitext(os.path.basename(path))[0]
                out_path = os.path.join(os.path.dirname(path), base + ".docx")

                self.ui_queue.put(("status", f"[{idx}/{total}] Transcrevendo: {os.path.basename(path)}"))

                ok = self._transcribe_one(path, out_path, idx, total)
                if ok:
                    last_out = out_path

            if self.cancel_requested:
                self.ui_queue.put(("canceled", "Transcrição cancelada."))
            else:
                if last_out:
                    self.last_output = last_out
                self.ui_queue.put(("done", f"Transcrição concluída de {total} arquivo(s)."))

        except Exception as e:
            if self.cancel_requested:
                self.ui_queue.put(("canceled", "Transcrição cancelada."))
            else:
                self.ui_queue.put(("error", f"Erro no processamento: {e}"))
        finally:
            self.is_running = False

    # Transcreve um único arquivo e salva .docx
    def _transcribe_one(self, in_path, out_docx, idx, total):
        try:
            # language=None => autodetect; para forçar pt-BR: language="pt"
            segments, info = self.model.transcribe(
                in_path,
                language=None,
                vad_filter=True,
                beam_size=5,
                condition_on_previous_text=True
            )

            # Junta o texto sem timestamps
            pieces = []
            for seg in segments:
                if self.cancel_requested:
                    break
                txt = (seg.text or "").strip()
                if txt:
                    pieces.append(txt)

            if self.cancel_requested:
                return False

            full_text = " ".join(pieces).strip()
            if not full_text:
                self.ui_queue.put(("error", f"Nenhum texto reconhecido em: {os.path.basename(in_path)}"))
                return False

            # Salva em DOCX
            try:
                doc = Document()
                doc.add_heading("Transcrição", level=1)
                if getattr(info, "language", None):
                    lang = info.language
                    doc.add_paragraph(f"Idioma detectado: {lang}")
                    doc.add_paragraph("")  # linha em branco
                doc.add_paragraph(full_text)
                doc.save(out_docx)
            except Exception as e:
                self.ui_queue.put(("error", f"Falha ao salvar DOCX ({os.path.basename(in_path)}): {e}"))
                return False

            self.ui_queue.put(("progress", 100))
            self.ui_queue.put(("status", f"[{idx}/{total}] Salvo: {os.path.basename(out_docx)}"))
            return True

        except Exception as e:
            self.ui_queue.put(("error", f"Erro ao transcrever {os.path.basename(in_path)}: {e}"))
            return False

    # ----------- UI Queue -----------
    def _drain_ui_queue(self):
        try:
            while True:
                kind, payload = self.ui_queue.get_nowait()
                if kind == "progress":
                    self.progress_var.set(payload)
                elif kind == "status":
                    self.status_var.set(payload)
                elif kind == "done":
                    self.progress_var.set(100)
                    self.status_var.set(payload)
                    self.btn_run.config(state=NORMAL)
                    self.btn_cancel.config(state=DISABLED)
                    self.btn_open.config(state=NORMAL if self.last_output else DISABLED)
                    messagebox.showinfo("Sucesso", "Transcrição concluída!")
                elif kind == "canceled":
                    self.status_var.set(payload)
                    self.progress_var.set(0)
                    self.btn_run.config(state=NORMAL)
                    self.btn_cancel.config(state=DISABLED)
                    self.btn_open.config(state=DISABLED)
                elif kind == "error":
                    messagebox.showerror("Erro", str(payload))
        except queue.Empty:
            pass
        finally:
            if not self.is_running and self.btn_cancel["state"] == NORMAL:
                self.btn_run.config(state=NORMAL)
                self.btn_cancel.config(state=DISABLED)
            self.after(100, self._drain_ui_queue)

    # ----------- Abrir pasta -----------
    def abrir_pasta(self):
        if self.last_output and os.path.exists(self.last_output):
            pasta = os.path.dirname(self.last_output)
            try:
                if sys.platform.startswith("win"):
                    os.startfile(pasta)
                elif sys.platform == "darwin":
                    subprocess.Popen(["open", pasta])
                else:
                    subprocess.Popen(["xdg-open", pasta])
            except Exception as e:
                messagebox.showerror("Erro", f"Não foi possível abrir a pasta: {e}")
        else:
            messagebox.showerror("Erro", "Nenhum arquivo transcrito encontrado.")


if __name__ == "__main__":
    app = AudioTranscriberApp()
    app.mainloop()
